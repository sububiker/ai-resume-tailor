"""Export TailoredResume to PDF (via HTML) and DOCX."""
import io

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib.enums import TA_CENTER, TA_LEFT

from models.schemas import TailoredResume


# ── PDF Export ────────────────────────────────────────────────────────────────

def export_pdf(resume: TailoredResume) -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )

    styles = getSampleStyleSheet()
    name_style = ParagraphStyle("Name", fontSize=20, fontName="Helvetica-Bold", alignment=TA_CENTER, spaceAfter=4)
    contact_style = ParagraphStyle("Contact", fontSize=9, fontName="Helvetica", alignment=TA_CENTER, textColor=colors.HexColor("#555555"), spaceAfter=2)
    section_style = ParagraphStyle("Section", fontSize=11, fontName="Helvetica-Bold", textColor=colors.HexColor("#1a1a2e"), spaceBefore=10, spaceAfter=4)
    body_style = ParagraphStyle("Body", fontSize=9.5, fontName="Helvetica", spaceAfter=2, leading=14)
    bullet_style = ParagraphStyle("Bullet", fontSize=9.5, fontName="Helvetica", leftIndent=12, spaceAfter=2, leading=14, bulletIndent=4)
    job_title_style = ParagraphStyle("JobTitle", fontSize=10, fontName="Helvetica-Bold", spaceAfter=1)
    sub_style = ParagraphStyle("Sub", fontSize=9, fontName="Helvetica", textColor=colors.HexColor("#555555"), spaceAfter=3)

    story = []
    pi = resume.personal_info

    # Name
    story.append(Paragraph(pi.name or "Your Name", name_style))

    # Contact line
    contact_parts = [x for x in [pi.email, pi.phone, pi.location, pi.linkedin] if x]
    if contact_parts:
        story.append(Paragraph(" · ".join(contact_parts), contact_style))

    story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor("#1a1a2e"), spaceAfter=6))

    # Summary
    if resume.summary:
        story.append(Paragraph("PROFESSIONAL SUMMARY", section_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#cccccc"), spaceAfter=4))
        story.append(Paragraph(resume.summary, body_style))

    # Experience
    if resume.experience:
        story.append(Paragraph("EXPERIENCE", section_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#cccccc"), spaceAfter=4))
        for exp in resume.experience:
            story.append(Paragraph(f"{exp.title} — {exp.company}", job_title_style))
            meta = " · ".join(x for x in [exp.duration, exp.location] if x)
            if meta:
                story.append(Paragraph(meta, sub_style))
            for bullet in exp.bullets:
                story.append(Paragraph(f"• {bullet}", bullet_style))
            story.append(Spacer(1, 4))

    # Education
    if resume.education:
        story.append(Paragraph("EDUCATION", section_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#cccccc"), spaceAfter=4))
        for edu in resume.education:
            story.append(Paragraph(edu.degree, job_title_style))
            meta = " · ".join(x for x in [edu.institution, edu.year, edu.gpa] if x)
            if meta:
                story.append(Paragraph(meta, sub_style))

    # Skills
    if resume.skills:
        story.append(Paragraph("SKILLS", section_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#cccccc"), spaceAfter=4))
        story.append(Paragraph(", ".join(resume.skills), body_style))

    # Certifications
    if resume.certifications:
        story.append(Paragraph("CERTIFICATIONS", section_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#cccccc"), spaceAfter=4))
        for cert in resume.certifications:
            story.append(Paragraph(f"• {cert}", bullet_style))

    doc.build(story)
    return buffer.getvalue()


# ── Word Export ───────────────────────────────────────────────────────────────

def _heading(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    doc.add_paragraph("─" * 60).paragraph_format.space_after = Pt(2)


def export_docx(resume: TailoredResume) -> bytes:
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Pt(54)
        section.bottom_margin = Pt(54)
        section.left_margin = Pt(54)
        section.right_margin = Pt(54)

    pi = resume.personal_info

    # Name
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run(pi.name or "Your Name")
    name_run.bold = True
    name_run.font.size = Pt(20)

    # Contact
    contact_parts = [x for x in [pi.email, pi.phone, pi.location, pi.linkedin] if x]
    if contact_parts:
        cp = doc.add_paragraph(" · ".join(contact_parts))
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cp.runs:
            run.font.size = Pt(9)

    # Summary
    if resume.summary:
        _heading(doc, "Professional Summary")
        doc.add_paragraph(resume.summary)

    # Experience
    if resume.experience:
        _heading(doc, "Experience")
        for exp in resume.experience:
            p = doc.add_paragraph()
            r = p.add_run(f"{exp.title} — {exp.company}")
            r.bold = True
            r.font.size = Pt(10)
            meta = " · ".join(x for x in [exp.duration, exp.location] if x)
            if meta:
                mp = doc.add_paragraph(meta)
                for run in mp.runs:
                    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                    run.font.size = Pt(9)
            for bullet in exp.bullets:
                doc.add_paragraph(bullet, style="List Bullet")

    # Education
    if resume.education:
        _heading(doc, "Education")
        for edu in resume.education:
            p = doc.add_paragraph()
            r = p.add_run(edu.degree)
            r.bold = True
            r.font.size = Pt(10)
            meta = " · ".join(x for x in [edu.institution, edu.year, edu.gpa] if x)
            if meta:
                doc.add_paragraph(meta)

    # Skills
    if resume.skills:
        _heading(doc, "Skills")
        doc.add_paragraph(", ".join(resume.skills))

    # Certifications
    if resume.certifications:
        _heading(doc, "Certifications")
        for cert in resume.certifications:
            doc.add_paragraph(cert, style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
